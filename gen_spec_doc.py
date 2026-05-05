# -*- coding: utf-8 -*-
"""
gen_spec_doc.py -- Generate asksancho-spec-v1.0.0.docx
Run: python3 gen_spec_doc.py
Output: dist/asksancho-spec-v1.0.0.docx
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = Path(__file__).parent / "dist"
OUTPUT_DIR.mkdir(exist_ok=True)
OUTPUT = OUTPUT_DIR / "req-refiner-spec-v1.0.0.docx"

# curly quotes as unicode escapes to avoid Python 3.12+ syntax restrictions
LQUOTE = "“"
RQUOTE = "”"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading1(doc, text):
    p = doc.add_heading(text, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        run.font.size = Pt(16)
    return p


def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x35, 0x74, 0x98)
        run.font.size = Pt(13)
    return p


def body(doc, text, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.6)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    run = p.add_run(text)
    set_font(run, size=11)
    return p


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_font(run, name="Courier New", size=9.5, color=(0x44, 0x44, 0x44))
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)
    return p


def table_2col(doc, rows, col_widths=(6.5, 9.5)):
    tbl = doc.add_table(rows=1 + len(rows), cols=2)
    tbl.style = "Table Grid"
    hdr_cells = tbl.rows[0].cells
    for i, label in enumerate(["项目 / Item", "内容 / Content"]):
        hdr_cells[i].text = label
        run = hdr_cells[i].paragraphs[0].runs[0]
        set_font(run, bold=True, size=10.5)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D0E4F7")
        tcPr.append(shd)
    for idx, (k, v) in enumerate(rows):
        cells = tbl.rows[idx + 1].cells
        cells[0].text = k
        set_font(cells[0].paragraphs[0].runs[0], bold=True, size=10)
        cells[1].text = v
        set_font(cells[1].paragraphs[0].runs[0], size=10)
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(int(col_widths[i] * 567)))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
    return tbl


def divider(doc):
    p = doc.add_paragraph()
    run = p.add_run("─" * 80)
    set_font(run, size=8, color=(0xCC, 0xCC, 0xCC))
    return p


# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Cover ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Requirement Refine — 开发规格说明书")
set_font(run, size=20, bold=True, color=(0x1A, 0x56, 0xDB))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("req-refiner-spec-v1.0.0  ·  2026-05-04")
set_font(run, size=11, color=(0x88, 0x88, 0x88))

doc.add_paragraph()
divider(doc)
doc.add_paragraph()


# ── Section 1: Why ───────────────────────────────────────────────────────────
heading1(doc, "1. 原因 / Why")

body(doc, (
    "用户用自然语言向 Claude Code 表达需求时，"
    "往往存在以下问题，"
    "导致 Claude Code 执行偏离预期、产生多轮返工："
))

bullet(doc, "需求边界模糊 — 没有明确 In Scope / Out of Scope，Claude 会自行扩展")
bullet(doc, "缺少成功标准 — 没有 Acceptance Criteria，无法判断" + LQUOTE + "什么是完成" + RQUOTE)
bullet(doc, "忽略已有约束 — CLAUDE.md / HANDOFF.md 中的规则未被纳入需求描述")
bullet(doc, "Opus vs Sonnet 使用方式不同 — Opus 需要留白供推理，Sonnet 需要无歧义指令")

doc.add_paragraph()
body(doc, (
    "本项目设计一套「需求精炼」 (Requirement Refine) 工作流："
    "在需求进入 Claude Code 之前，"
    "通过结构化对话完成澄清、补充、限定范围，"
    "最终为不同模型生产高质量提示词。"
))
doc.add_paragraph()


# ── Section 2: Original Instructions ─────────────────────────────────────────
heading1(doc, "2. 原始开发指令 / Original Instructions")

body(doc, "用户在 Claude Code 中输入（2026-04-28，会话cdb9363f）：")
doc.add_paragraph()

original_cn = (
    "规划一组整理自然语言输入 Claude Code 的工作流，"
    "在需求进入 Claude Code 前先进行澄清、结构化，"
    "最终对 Opus 4.7 和 Sonnet 4.6 分别给出高质量的提示词 / 需求文档。三层架构：\n"
    "· 基础版 — 独立提示词，任何 LLM 对话窗口可用，零安装\n"
    "· 进阶版 — Claude Code skill，在项目上下文中运行，产出 plan 文件\n"
    "· 高级版 — 本地 LLM 中间层（Gemma 4 + 向量记忆），支持多模态输入，结构化需求直写 Claude Code"
)
code_block(doc, original_cn)

doc.add_paragraph()
body(doc, "用户补充指令（ultraplan 结果整合阶段）：")
doc.add_paragraph()

original_en = (
    "吸收整合 ultraplan 生成的规划，做一个最终的完整方案，"
    "并生成一个 word 版本的开发说明，包括原因、原始的开发指令，"
    "解析后的需求，实现策略和最终的产出方案描述等。"
)
code_block(doc, original_en)
doc.add_paragraph()


# ── Section 3: Parsed Requirements ───────────────────────────────────────────
heading1(doc, "3. 解析后的需求 / Parsed Requirements")

heading2(doc, "3.1 Goal")
body(doc, (
    "在用户的自然语言需求进入 Claude Code 之前，"
    "通过一套分层工具链对需求进行澄清和结构化，"
    "使 Claude Code（Opus 4.7 / Sonnet 4.6）收到的提示词足够清晰、边界明确、可验证。"
))

heading2(doc, "3.2 In Scope")
bullet(doc, "基础版：独立 meta-prompt，任何 LLM 可用，5 步对话协议，双模型输出")
bullet(doc, "进阶版：Claude Code skill（/requirement-refine），"
             "读取 CLAUDE.md/HANDOFF.md/git log，"
             "产出结构化 spec，支持 3 种移交方式（plan mode / output only / scratch 文件桥）")
bullet(doc, "高级版：架构探索文档（非立即实现），涵盖本地 LLM 选型、向量 DB、多模态输入、Bridge 选型")
bullet(doc, "Word 版开发说明文档（本文件）")

heading2(doc, "3.3 Out of Scope")
bullet(doc, "高级版的实际代码实现（除 collector.py 骨架外）")
bullet(doc, "多语言 i18n 支持")
bullet(doc, "与外部 CI/CD 系统集成")
bullet(doc, "与 AIGP 项目代码库耦合（本项目为通用独立 repo）")

heading2(doc, "3.4 Inputs / Outputs")
rows_io = [
    ("输入", "用户自然语言需求（文本；高级版额外支持语音/图像）"),
    ("输出 — 基础版", "结构化 spec（7 字段模板），含 Opus 4.7 + Sonnet 4.6 差异化提示"),
    ("输出 — 进阶版", "结构化 spec（7 字段 + Code references），移交给 plan mode 或 scratch 文件"),
    ("输出 — 高级版", "架构文档（组件选型 + Bridge 路线图），为 Phase P1–P5 提供决策依据"),
]
table_2col(doc, rows_io)

heading2(doc, "3.5 Constraints")
bullet(doc, "进阶版 skill 只读取轻量上下文（CLAUDE.md head -80 / HANDOFF.md head -60 / git log -8），不 glob 代码库")
bullet(doc, "澄清轮次上限：3 轮，每轮最多 4 个问题（AskUserQuestion 上限）")
bullet(doc, "不提实现建议 — skill 只澄清" + LQUOTE + "做什么" + RQUOTE + "，不讨论" + LQUOTE + "怎么做" + RQUOTE)
bullet(doc, "高级版本地 LLM：Gemma 4 via Ollama，向量 DB：ChromaDB（零配置）")
bullet(doc, "独立 repo（不绑定 AIGP），通用适用")

heading2(doc, "3.6 Acceptance Criteria")
bullet(doc, "[ ] 基础版 meta-prompt 能在 Claude.ai / ChatGPT 中完成 5 步协议并输出合规 spec")
bullet(doc, "[ ] 进阶版 /requirement-refine skill 在 AIGP 项目中读取 CLAUDE.md 规则并纳入 spec")
bullet(doc, "[ ] 进阶版 spec 的 Code references 节来自实际 git/文件上下文，非猜测")
bullet(doc, "[ ] 高级版架构文档覆盖 Bridge 4 个候选方案及推荐路线图")
bullet(doc, "[ ] Word 文档包含原因/指令/需求/策略/方案描述 5 个章节，可独立阅读")

doc.add_paragraph()


# ── Section 4: Implementation Strategy ───────────────────────────────────────
heading1(doc, "4. 实现策略 / Implementation Strategy")

heading2(doc, "4.1 分层交付策略")
body(doc, (
    "三层架构从零依赖到深度集成逐层递进，"
    "每层独立可用，不阻塞下层实现。"
    "本次会话完成基础版 + 进阶版，高级版以架构文档形式交付待后续实现。"
))

heading2(doc, "4.2 五步对话协议（三层共用）")
rows_5step = [
    ("Step 1 — Restate",
     "一句话复述需求，以 \"What I read is: ... Is that right?\" 收尾"),
    ("Step 2 — Disambiguate",
     "识别 ≤5 条歧义，每次只问一条，等回答再问下一条"),
    ("Step 3 — Probe",
     "主动追问 4 类遗漏：成功标准 / Out of Scope / 边界条件 / 失败行为"),
    ("Step 4 — Consolidate",
     "汇总已澄清内容，问“还有遗漏的约束吗？”等确认"),
    ("Step 5 — Output",
     "填入 7 字段 spec 模板，不加任何寒教或总结"),
]
table_2col(doc, rows_5step, col_widths=(5.0, 11.0))

heading2(doc, "4.3 Opus 4.7 vs Sonnet 4.6 差异化策略")
rows_model = [
    ("Spec 风格",
     "Opus：保留模糊性，含" + LQUOTE + "为什么" + RQUOTE + "背景  ·  Sonnet：消除歧义，有序任务列表"),
    ("Spec 末尾",
     "Opus：\"Please plan before editing.\"  ·  Sonnet：具体验证命令（如 python verify.sh）"),
    ("文件路径",
     "Opus：可不写，模型会探索  ·  Sonnet：必须写清楚已知路径和函数名"),
    ("替代方案",
     "Opus：可邀请替代方案  ·  Sonnet：不邀请，指定做法"),
]
table_2col(doc, rows_model, col_widths=(4.0, 12.0))

heading2(doc, "4.4 进阶版 Bridge 移交（3 选项）")
rows_bridge = [
    ("A — Plan mode", "将 spec 作为 plan mode 起点，执行 EnterPlanMode"),
    ("B — Output only", "原样输出 spec，用户自行复制给 Claude Code"),
    ("C — Scratch file", "写入 ~/.claude/scratch/last-requirement-spec.md，跨会话可引用"),
]
table_2col(doc, rows_bridge, col_widths=(4.5, 11.5))

heading2(doc, "4.5 高级版 Bridge 路线图（架构决策）")
rows_bridge2 = [
    ("P1 (MVP)", "方案 A：剪贴板（pyperclip），零依赖，手动粘贴"),
    ("P2", "方案 B：文件桥（~/.claude/scratch/），已在进阶版 Option C 实现"),
    ("P3 (正式版)", "方案 C：MCP Server（JSON-RPC），深度 Claude Code 集成"),
    ("P4 (探索)", "方案 D：OS-level 注入（AppleScript），需 Accessibility 权限，慎用"),
]
table_2col(doc, rows_bridge2, col_widths=(3.5, 12.5))

doc.add_paragraph()


# ── Section 5: Final Deliverables ────────────────────────────────────────────
heading1(doc, "5. 最终产出方案描述 / Final Deliverables")

heading2(doc, "5.1 基础版 — 独立 Meta-Prompt")
rows_basic = [
    ("文件", "~/.claude/skills/requirement-refine/prompts/basic-v1.0.en.md（英文部署版）\n"
              "~/.claude/skills/requirement-refine/prompts/basic-v1.0.bilingual.md（中英审阅版）"),
    ("用途", "用户复制全文 → 粘贴到任意 LLM 对话窗口 → 在文末写需求 → 开始对话"),
    ("协议", "五步对话协议（Restate → Disambiguate → Probe → Consolidate → Output）"),
    ("硬性约束", "不读取代码文件 · 一次只问一条歧义 · 不提实现建议"),
    ("输出", "7 字段 spec 模板 + Opus 4.7 / Sonnet 4.6 差异化 hints"),
    ("状态", "✅ 可用"),
]
table_2col(doc, rows_basic)

doc.add_paragraph()
heading2(doc, "5.2 进阶版 — Claude Code Skill")
rows_adv = [
    ("Skill 名", "requirement-refine（触发：/requirement-refine [需求]）"),
    ("文件", "~/.claude/skills/requirement-refine/SKILL.md"),
    ("上下文读取", "CLAUDE.md head -80 · HANDOFF.md head -60 · git log --oneline -8"),
    ("澄清", "AskUserQuestion，≤3 轮，每轮 ≤4 问，按影响大小排序"),
    ("Spec 额外字段", "Code references（来自实际项目上下文，不猜测）"),
    ("移交", "A=plan mode · B=output only · C=~/.claude/scratch/last-requirement-spec.md"),
    ("状态", "✅ 可用"),
]
table_2col(doc, rows_adv)

doc.add_paragraph()
heading2(doc, "5.3 高级版 — 架构文档（未实现）")
rows_prem = [
    ("架构文档", "~/.claude/skills/requirement-refine/docs/advanced-tier-architecture.md"),
    ("本地 LLM", "Gemma 4 via Ollama（推荐）"),
    ("向量 DB", "ChromaDB（零配置，本地持久化）"),
    ("语音", "whisper.cpp（本地，MLX 加速，Mac M-series）"),
    ("Web UI", "FastAPI + 简单前端（app.py + ui/index.html）"),
    ("Bridge 推荐", "P1 剪贴板 → P2 文件桥 → P3 MCP Server"),
    ("实现分阶段", "P1 记忆层 → P2 精炼器 → P3 Web UI → P4 多模态 → P5 MCP 注入"),
    ("状态", "📋 架构文档阶段（core/ 骨架已创建，逻辑待实现）"),
]
table_2col(doc, rows_prem)

doc.add_paragraph()
heading2(doc, "5.4 独立 Repo 结构")
code_block(doc, """\
req-refiner/
├── README.md
├── dist/
│   └── req-refiner-spec-v1.0.0.docx   ← 本文件
├── prompts/
│   └── req-refiner-basic.md            ← 基础版 meta-prompt
├── examples/
│   └── example-session.md              ← 示例对话
├── skill/
│   └── SKILL.md                        ← 进阶版 skill
├── core/
│   ├── collector.py                    ← 已实现
│   ├── indexer.py                      ← TODO Phase A
│   ├── refiner.py                      ← TODO Phase B
│   └── injector.py                     ← TODO Phase E
├── ui/
│   └── index.html                      ← TODO Phase C
├── app.py                              ← TODO Phase C
└── requirements.txt

~/.claude/skills/requirement-refine/    ← 已部署
├── SKILL.md                            ← 进阶版主文件
├── prompts/
│   ├── basic-v1.0.en.md
│   ├── basic-v1.0.bilingual.md
│   └── README.md                       ← 三层选用决策表
└── docs/
    └── advanced-tier-architecture.md   ← 高级版架构探索文档""")

doc.add_paragraph()
divider(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("req-refiner-spec-v1.0.0  ·  Generated 2026-05-04")
set_font(run, size=9, color=(0xAA, 0xAA, 0xAA))

doc.save(str(OUTPUT))
print(f"Saved: {OUTPUT}")
